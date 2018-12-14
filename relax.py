#!/usr/bin/env python

import sys
import nmap

from docx import Document
from docx.shared import Inches

VER = 2

try:
    if sys.version_info >= (3, 0):
        VER = 3
        from urllib.request import urlopen
        from urllib.error import URLError
        raw_input = input
    else:
        from urllib2 import urlopen
        from urllib2 import URLError
except:
        pass

documento = Document()

def validarespuesta(sitios):
    count = 0
    for sitio in sitios:
        count = count + 1
        print str(count) + ': ' + sitio
    respuesta = raw_input('\n' +
                          '\033[91mEs correcta la lista de los sitios? (S/N): \033[0m')
    if respuesta.lower() == 's':
        comenzar_escaneo(sitios)
    elif respuesta.lower() == 'n':
        index = raw_input('\033[91mQue sitio esta'
                          ' mal? Introduce el numero identificador:\033[0m ')
        sitios.pop(int(index)-1)
        validarespuesta(sitios)
    else:
        print 'Escriba un comando correcto por favor'
        validarespuesta(sitios)

def fetch(url, decoding='utf-8'):
    "Obtiene el contenido de la url"
    return urlopen(url).read().decode(decoding)

def scannmap(host):
    "Realiza el escaneo del host con nmap"
    print '\033[92m[+]\033[0mComenzando escaneo con Nmap de: ' + host + '\033[92m[+]\033[0m'
    scan = nmap.PortScanner()
    scan.scan(host)
    for host in scan.all_hosts():
        print '----------------------------------------------------'
        print 'Host : %s (%s)' % (host, scan[host].hostname())
        print 'Estado : %s' % scan[host].state()
        for proto in scan[host].all_protocols():
            print '----------'
            print 'Protocolo : %s' % proto
            lport = scan[host][proto].keys()
            lport.sort()
            for port in lport:
                documento.add_paragraph('Puerto : %s\tEstado : %s\tUtilizado para: %state - %s' % (port, scan[host][proto][port]['state'], scan[host][proto][port]['product'], scan[host][proto][port]['version']))

def whois(host):
    "Realiza un whois al host"
    print '\n\033[92m[+]\033[0mComenzando escaneo con WHOIS de: ' + host + '\033[92m[+] \033[0m'
    url = "http://api.hackertarget.com/whois/?q=" + host
    pwho = fetch(url)
    documento.add_paragraph(pwho)

def dnslookup(host):
    print '\n\033[92m[+]\033[0mComenzando escaneo con DNSLOOKUP de: ' + host + '\033[92m[+] \033[0m'
    "Realiza un dnslookup al host"
    ns = "http://api.hackertarget.com/dnslookup/?q=" + host
    pns = fetch(ns)
    documento.add_paragraph(pns)

def page_links(host):
    print '\n\033[92m[+]\033[0mComenzando a escanear las URL de la pagina: ' + host + '\033[92m[+] \033[0m'
    ns = "https://api.hackertarget.com/pagelinks/?q=" + host
    res = fetch(ns)
    documento.add_paragraph(res)

def test_ping(host):
    print '\n\033[92m[+]\033[0mComenzando el testeo de los ping de: ' + host + '\033[92m[+] \033[0m'
    ns = "https://api.hackertarget.com/nping/?q=" + host
    res = fetch(ns)
    documento.add_paragraph(res)

def tracerouter(host):
    print '\n\033[92m[+]\033[0mRealizando tracer route sobre: ' + host + '\033[92m[+] \033[0m'
    ns = "https://api.hackertarget.com/mtr/?q=" + host
    res = fetch(ns)
    documento.add_paragraph(res)


def comenzar_escaneo(sitios):
    "Comienza el escaneo completo de los sitios ingresados"
    for host in sitios:
        documento.add_heading('Analisis sobre ' + host, 0)
        
        documento.add_heading('Parte legal:',level=2)
        documento.add_paragraph('Todos los derechos reservados').bold = True
        documento.add_paragraph(
        'Este documento contiene informacion confidencial. Este documento es de uso ' +
        'exclusivo para la empresa ' + host + '. El uso inautorizado o reproduccion de este ' +
        'documento esta estrictamente prohibido. Este pentesting ha sido manejado por ' +
        'expertos en seguridad. El equipo asegura que las ' +
        'vulnerabilidades encontradas y escritas en este reportes son totalmente ' +
        'verdaderas y se pueden verificar via internet. Este reporte de pentesting muestra ' +
        'todas las vulnerabilidades conocidas de '+host+' a la fecha del analisis. A medida que ' +
        'nuevas vulneravilidades sean liberadas o encontrar posibles amenazas a la ' +
        'seguridad del sistema es sugerido que las asesorias de seguridad y las pruebas ' +
        'de pentesting se realizen dentro de un intervalo de 3 a 6 meses.')
        
        documento.add_heading('Detalles del documento', level=2)
        table = documento.add_table(rows = 1, cols = 2)
        row_cells = table.add_row().cells
        row_cells[0].text = 'Tipo de documento'
        row_cells[1].text = 'Reporte de analisis de vulnerabilidades'

        documento.add_heading('Resumen ejecutivo', level=2)
        documento.add_paragraph(host + ' contrato a un equipo de seguridad con el fin de realizar un test de pentesting ' +
            'cuyo objetivo es determinar su debilidad en contra de un intento de ataque ' +
            'informatico. Todas las actividades fueron manejadas de manera que se simulo ' +
            'un ataque a cargo de un actor maligno en contra de la aplicacion de ' + host)

        documento.add_paragraph('La simulacion sobre el ataque unicamente llego hasta la etapa de analisis de ' +
                'informacion, los ataques y analisis encontrados durante el test de penetracion ' +
                'se encontraron con el nivel de acceso con el que cualquier usuario general de ' +
                'internet pudiera contar')

        documento.add_heading("Resultados del analisis con NMAP", level=1)
        scannmap(host)
        documento.add_heading("Resultados del analisis con WHOIS", level=1)
        whois(host)
        documento.add_heading("Resultados del analisis con DNSLOOKUP", level=1)
        dnslookup(host)
        documento.add_heading("Resultados del analisis con TRACEROUTER", level=1)
        tracerouter(host)
        documento.add_heading("Resultados del analisis con TEST PING", level=1)
        test_ping(host)
        documento.add_heading("Resultados del analisis con PAGE LINKS", level=1)
        page_links(host)
        documento.save('Analisis - '+ host + '.docx')

def banner():
    print '\033[92m/$$$$       /$$$$\033[0m       /$$$$$$$            | $$                          \033[92m /$$$$       /$$$$\033[0m'
    print '\033[92m| $$_/  /$$ |_  $$\033[0m      | $$__  $$          | $$                          \033[92m| $$_/  /$$ |_  $$\033[0m'
    print '\033[92m| $$   | $$   | $$\033[0m      | $$  \ $$  /$$$$$$ | $$  /$$$$$$  /$$   /$$      \033[92m| $$   | $$   | $$\033[0m'
    print '\033[92m| $$ /$$$$$$$$| $$\033[0m      | $$$$$$$/ /$$__  $$| $$ |____  $$|  $$ /$$/      \033[92m| $$ /$$$$$$$$| $$\033[0m'
    print '\033[92m| $$|__  $$__/| $$\033[0m      | $$__  $$| $$$$$$$$| $$  /$$$$$$$ \  $$$$/       \033[92m| $$|__  $$__/| $$\033[0m'
    print '\033[92m| $$   | $$   | $$\033[0m      | $$  \ $$| $$_____/| $$ /$$__  $$  >$$  $$       \033[92m| $$   | $$   | $$\033[0m'
    print '\033[92m| $$$$ |__/  /$$$$\033[0m      | $$  | $$|  $$$$$$$| $$|  $$$$$$$ /$$/\  $$      \033[92m| $$$$ |__/  /$$$$\033[0m'
    print '\033[92m|____/      |____/\033[0m      |__/  |__/ \_______/|__/ \_______/|__/  \__/      \033[92m|____/      |____/\033[0m'
    print '\t\t\t\tHecho con mucho \033[91m<3\033[0m por Hector Camacho!'
    print '\n'
    menu()

def inicio():
    banner()
    sitios = []
    sitio = raw_input('\033[91mIntroduce algun sitio '
                      'para analizar:\033[0m ')
    sitios.append(sitio)
    while sitio != 'Listo' and sitio != 'Salir':
        sitio = raw_input('\033[91mIntroduce algun '
                          'sitio para analizar:\033[0m ')
        if sitio == 'Listo':
            validarespuesta(sitios)
            break;
        if sitio == 'Salir':
            return 'Saliendo...';
            break;
        else:
            sitios.append(sitio)

def menu():
    print '\n'
    print 'Una vez que los sitios esten capturados escribe "Listo" para continuar con el escaneo.\n'
    print 'Para salir del programa escribe "Salir"\n'

if __name__ == '__main__':
    inicio()
